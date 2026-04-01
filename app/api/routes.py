import logging
import uuid
from pathlib import Path

from fastapi import APIRouter, BackgroundTasks, HTTPException, UploadFile
from fastapi.responses import FileResponse
from typing import List

from app.core import extractor, parser, reporter, zip_extractor
from app.core.config import settings
from app.core.database import get_session
from app.models.db import PolicyRecord, Task

logger = logging.getLogger(__name__)

router = APIRouter()

UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")

# 支持直接上传的单文件格式（不含 zip，zip 单独处理）
SINGLE_FILE_EXTENSIONS = {".pdf", ".docx", ".doc", ".wps", ".jpg", ".jpeg", ".png", ".bmp", ".tiff"}


@router.post("/upload", summary="批量上传文件（.zip / .pdf / .docx / 图片），所有文件合并为一个任务，导出一份 Excel")
async def upload_files(background_tasks: BackgroundTasks, files: List[UploadFile]):
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")

    UPLOAD_DIR.mkdir(exist_ok=True)

    task_id = str(uuid.uuid4())
    zip_paths: list[Path] = []
    single_paths: list[Path] = []
    zip_idx = 0
    single_idx = 0

    for file in files:
        if not file.filename:
            continue
        suffix = Path(file.filename).suffix.lower()
        content = await file.read()

        if suffix == ".zip":
            dest = UPLOAD_DIR / f"{task_id}_z{zip_idx}.zip"
            dest.write_bytes(content)
            zip_paths.append(dest)
            zip_idx += 1
        elif suffix in SINGLE_FILE_EXTENSIONS:
            dest = UPLOAD_DIR / f"{task_id}_s{single_idx}{suffix}"
            dest.write_bytes(content)
            single_paths.append(dest)
            single_idx += 1
        else:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件格式：'{file.filename}'。支持 .zip / .pdf / .docx / .doc / .wps / .jpg / .jpeg / .png / .bmp / .tiff",
            )

    if not zip_paths and not single_paths:
        raise HTTPException(status_code=400, detail="未找到有效文件")

    filenames = ",".join(f.filename for f in files if f.filename)
    with get_session() as db:
        db.add(Task(id=task_id, status="pending", filename=filenames))
        db.commit()

    background_tasks.add_task(_process_task, task_id, zip_paths, single_paths)
    return {"task_id": task_id, "filenames": [f.filename for f in files if f.filename], "status": "pending"}


@router.get("/tasks/{task_id}", summary="查询任务状态")
def get_task(task_id: str):
    with get_session() as db:
        task = db.get(Task, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    return {"task_id": task.id, "status": task.status, "error": task.error_msg}


@router.get("/tasks/{task_id}/download", summary="下载本次任务的 Excel 报告")
def download_result(task_id: str):
    with get_session() as db:
        task = db.get(Task, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if task.status != "done":
        raise HTTPException(status_code=400, detail=f"Task is not ready (status: {task.status})")

    output_path = OUTPUT_DIR / f"{task_id}.xlsx"
    if not output_path.exists():
        raise HTTPException(status_code=404, detail="Output file not found")

    return FileResponse(
        path=str(output_path),
        filename=f"policy_summary_{task_id[:8]}.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def _extract_source_url_from_files(files: list[Path]) -> str | None:
    """从解压后的文件列表中提取来源 URL。

    按文件名排序，依次处理 .docx/.doc/.wps/.pdf 文件（图片跳过），
    读取每个文件的第一个非空段落/行，用正则匹配 URL，
    返回找到的第一个 URL；找不到则返回 None。
    """
    import re

    _IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff"}

    sorted_files = sorted(files, key=lambda p: p.name)
    for file_path in sorted_files:
        if file_path.name.startswith("~$"):
            continue
        suffix = file_path.suffix.lower()
        if suffix in _IMAGE_EXTS:
            continue

        try:
            first_line: str | None = None

            if suffix == ".docx":
                import io
                import docx as _docx
                doc = _docx.Document(io.BytesIO(file_path.read_bytes()))
                for para in doc.paragraphs:
                    t = para.text.strip()
                    if t:
                        first_line = t
                        break

            elif suffix in (".doc", ".wps"):
                # 先尝试 python-docx（部分 .doc/.wps 实际是 OOXML）
                try:
                    import io
                    import docx as _docx
                    doc = _docx.Document(io.BytesIO(file_path.read_bytes()))
                    for para in doc.paragraphs:
                        t = para.text.strip()
                        if t:
                            first_line = t
                            break
                except Exception:
                    # 降级到 win32com COM 自动化
                    try:
                        import win32com.client  # type: ignore[import]
                        _progids = (
                            ("WPS.Application", "KWPS.Application", "Word.Application")
                            if suffix == ".wps"
                            else ("Word.Application",)
                        )
                        for progid in _progids:
                            app = None
                            try:
                                app = win32com.client.Dispatch(progid)
                                app.Visible = False
                                doc_obj = app.Documents.Open(str(file_path.absolute()))
                                all_text = doc_obj.Content.Text
                                doc_obj.Close(False)
                                for line in all_text.splitlines():
                                    t = line.strip()
                                    if t:
                                        first_line = t
                                        break
                            except Exception:
                                pass
                            finally:
                                if app is not None:
                                    try:
                                        app.Quit()
                                    except Exception:
                                        pass
                            if first_line:
                                break
                    except ImportError:
                        pass

            elif suffix == ".pdf":
                import io
                import fitz  # PyMuPDF
                doc = fitz.open(stream=file_path.read_bytes(), filetype="pdf")
                if doc.page_count > 0:
                    page_text = doc[0].get_text()
                    lines = [ln.strip() for ln in page_text.splitlines() if ln.strip()]
                    first_line = "\n".join(lines[:10]) if lines else None
                doc.close()

            if first_line:
                match = re.search(r"https?://\S+", first_line)
                if match:
                    return match.group(0)

        except Exception:
            continue

    return None


def _process_task(task_id: str, zip_paths: list[Path], single_paths: list[Path]) -> None:
    """后台任务：处理 ZIP 和单文件 → LLM 提取 → 写库 → 导出一份合并 Excel。

    每个 ZIP 内的所有文档合并为一次 LLM 提取；每个单文件独立进行一次 LLM 提取。
    """

    def _set_status(status: str, error_msg: str | None = None) -> None:
        with get_session() as db:
            task = db.get(Task, task_id)
            if task:
                task.status = status
                task.error_msg = error_msg
                db.commit()

    _set_status("processing")

    try:
        all_rows = []

        # ── ZIP 文件：解压 → 合并文本 → 一次 LLM 提取 ──
        for idx, zip_path in enumerate(zip_paths):
            extract_dir = UPLOAD_DIR / f"{task_id}_z{idx}"
            logger.info("Task %s ZIP[%d/%d] 开始解压: %s", task_id, idx + 1, len(zip_paths), zip_path.name)
            files = zip_extractor.safe_extract(zip_path, extract_dir)
            logger.info("Task %s ZIP[%d] 解压完成，共 %d 个文件", task_id, idx + 1, len(files))

            texts = []
            for file_path in files:
                logger.info("Task %s 解析文件: %s", task_id, file_path.name)
                text = parser.parse_document(
                    file_path,
                    ocr_api_url=settings.paddle_ocr_api_url,
                    ocr_api_key=settings.paddle_ocr_api_key,
                )
                logger.info("Task %s 文件解析完成，文本长度 %d 字符", task_id, len(text))
                texts.append(text)

            merged_text = "\n\n---\n\n".join(texts) if len(texts) > 1 else texts[0]
            logger.info("Task %s ZIP[%d] 开始 LLM 提取，合并文本 %d 字符", task_id, idx + 1, len(merged_text))
            rows = extractor.extract(merged_text)
            logger.info("Task %s ZIP[%d] LLM 提取完成，共 %d 行", task_id, idx + 1, len(rows))

            source_url = _extract_source_url_from_files(files)
            if source_url:
                logger.info("Task %s ZIP[%d] 从文档首行提取到来源链接，强制覆盖网站链接: %s", task_id, idx + 1, source_url)
                for row in rows:
                    row.网站链接 = source_url

            all_rows.extend(rows)

        # ── 单文件：直接解析 → 独立一次 LLM 提取 ──
        for idx, file_path in enumerate(single_paths):
            logger.info("Task %s 单文件[%d/%d] 解析: %s", task_id, idx + 1, len(single_paths), file_path.name)
            text = parser.parse_document(
                file_path,
                ocr_api_url=settings.paddle_ocr_api_url,
                ocr_api_key=settings.paddle_ocr_api_key,
            )
            logger.info("Task %s 单文件[%d] 解析完成，文本长度 %d 字符", task_id, idx + 1, len(text))
            logger.info("Task %s 单文件[%d] 开始 LLM 提取", task_id, idx + 1)
            rows = extractor.extract(text)
            logger.info("Task %s 单文件[%d] LLM 提取完成，共 %d 行", task_id, idx + 1, len(rows))

            source_url = _extract_source_url_from_files([file_path])
            if source_url:
                logger.info("Task %s 单文件[%d] 从文档首行提取到来源链接，强制覆盖网站链接: %s", task_id, idx + 1, source_url)
                for row in rows:
                    row.网站链接 = source_url

            all_rows.extend(rows)

        # ── 写入 MySQL ──
        with get_session() as db:
            for row in all_rows:
                db.add(PolicyRecord(task_id=task_id, **row.model_dump()))
            db.commit()
            records = db.query(PolicyRecord).filter(PolicyRecord.task_id == task_id).all()

        # ── 导出 Excel ──
        reporter.export_to_excel(task_id, records, OUTPUT_DIR)

        _set_status("done")

    except Exception as e:
        logger.exception("Task %s failed", task_id)
        _set_status("error", error_msg=str(e))
